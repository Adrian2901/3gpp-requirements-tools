import os
import requests
import json
import cv2
import easyocr
import base64
from docx import Document
from docx.shared import Inches
from docx2python import docx2python
from PIL import Image
from io import BytesIO, StringIO

def preprocess_image(image_path):
    '''
    Preprocess the image by applying different kernels to make it easier for OCR to detect text.
    '''
    # Load the image
    img = cv2.imread(image_path)

    # Convert to grayscale
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)

    # Apply thresholding to make the text stand out
    _, binary = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)

    # Kernels for removing horizontal and vertical lines
    kernel_h = cv2.getStructuringElement(cv2.MORPH_RECT, (50, 1))
    kernel_v = cv2.getStructuringElement(cv2.MORPH_RECT, (1, 50))

    # Remove horizontal lines
    horizontal_lines = cv2.morphologyEx(binary, cv2.MORPH_OPEN, kernel_h, iterations=2)
    binary_no_hlines = cv2.subtract(binary, horizontal_lines)

    # Remove vertical lines
    vertical_lines = cv2.morphologyEx(binary_no_hlines, cv2.MORPH_OPEN, kernel_v, iterations=1)
    binary_no_lines = cv2.subtract(binary_no_hlines, vertical_lines)
    
    # Apply slight blur to reduce noise
    blurred_img = cv2.medianBlur(binary_no_lines, 1)

    # Invert the binary image
    preprocessed_img = cv2.bitwise_not(blurred_img)

    return preprocessed_img


def process_sequence_diagram(image_path, debug=False):
    ''' 
    Process a sequence diagram image to extract actors and messages.
    '''
    # Threshold for actors detection (anything above this threshold is considered an actor)
    actors_threshold = 0.08

    # Preprocess the image
    img = preprocess_image(image_path)
    # DEBUG: Save the preprocessed image
    if debug:
        cv2.imwrite("preprocessed_image.png", img)

    # Use the easyocr model to extract text
    reader = easyocr.Reader(['en'])
    results = reader.readtext(img)

    # Lists for storing relevant texts and bounding boxes
    texts, bboxes = [], []
    # Get text and bounding box info from easyocr results
    for (bbox, text, conf) in results:
        # Remove empty strings
        if text.strip(): 
            # Remove the "|" character (sometimes dashed lines are detected by the ocr)
            if text != "|":
                # Add the text and bounding box coordinates to the lists
                texts.append(text)
                x_min = min(bbox, key=lambda p: p[0])[0]
                y_min = min(bbox, key=lambda p: p[1])[1]
                x_max = max(bbox, key=lambda p: p[0])[0]
                y_max = max(bbox, key=lambda p: p[1])[1]
                bboxes.append((x_min, y_min, x_max - x_min, y_max - y_min))

    # Lists for storing combined overlapping texts
    combined_texts = []
    combined_bboxes = []

    # Iterate over the texts and bboxes to combine overlapping texts
    for i, (text, bbox) in enumerate(zip(texts, bboxes)):
        # Get the bounding box coordinates
        x, y, w, h = bbox
        # If the text is above the actors threshold or there are no combined bboxes yet
        if not combined_bboxes or y < img.shape[0] * actors_threshold:
            combined_texts.append(text)
            combined_bboxes.append(bbox)
        else:
            # Get the previous bbox
            prev_x, prev_y, prev_w, prev_h = combined_bboxes[-1]

            # Check if the current bbox overlaps with the previous one in height
            overlap_height = min(y + h, prev_y + prev_h) - max(y, prev_y)
            # If the overlap is more than half of the smaller height, combine the texts
            if overlap_height >= 0.5 * min(h, prev_h):
                # Combine the texts
                combined_texts[-1] += " " + text

                # Combine the bboxes
                new_x = min(prev_x, x)
                new_y = min(prev_y, y)
                new_w = max(prev_x + prev_w, x + w) - new_x
                new_h = max(prev_y + prev_h, y + h) - new_y
                combined_bboxes[-1] = (new_x, new_y, new_w, new_h)
            else:
                # If there is no overlap, add the text and bbox to the lists
                combined_texts.append(text)
                combined_bboxes.append(bbox)

    # Update the texts and bboxes lists
    texts = combined_texts
    bboxes = combined_bboxes

    # DEBUG: Save an image with bounding boxes drawn over the detected texts
    if debug:
        for (x, y, w, h) in bboxes:
            cv2.rectangle(img, (int(x), int(y)), (int(x + w), int(y + h)), (0, 255, 0), 2)
        cv2.imwrite("debug_image.png", img)

    # Identify actors and messages
    actors = []
    messages = []
    # Iterate over the texts and bboxes
    for text, (x, y, w, h) in zip(texts, bboxes):
        # Check if the text is above the actors threshold
        if y < img.shape[0] * actors_threshold:
            # Store actor name and x-center (to be used in arrow detection)
            actors.append((text, x + w // 2)) 
        else:
            # Store message and bbox
            messages.append((text, (x, y, w, h)))  

    # Output string
    output = ""
    # Print actors and messages
    output += "Actors:\n"
    for actor, _ in actors:
        output += f"\t{actor}\n"
    output += "Messages:\n"
    for message in messages:
        output += f"\t{message}\n"

    return output

def is_sequence_diagram(image_path, llm_address, prompts):
    '''
    Ask the multimodal LLM whether the attached image is a sequence diagram.
    :param image_path: The path to the image to be analyzed
    :return: Boolean value indicating whether the image is a sequence diagram
    '''
    
    # Encode the image to base64 so that it can be sent to the LLM model
    encoded_image = base64.b64encode(open(image_path, "rb").read()).decode('utf-8')

    # Construct the data to be sent to the LLM model
    prompt_text = prompts['verify_image_context']
    url = f'http://{llm_address}/api/generate'
    data = {
        "model": "minicpm-v",
        "prompt": prompt_text,
        "images": [encoded_image],
        "stream": False
        }
    headers = {'Content-Type': 'application/json'}

    # Send the request to the LLM model and extract the response from the JSON data
    try:
        response = requests.post(url, data=json.dumps(data), headers=headers)
        json_data = json.loads(response.text)
        text = json_data['response']
    except requests.exceptions.RequestException as e:
        print(f"Error prompting the LLM: {e}")
        # If the request doesn't go through, return True, to not lose anything
        return True

    # List storing "relevant" answers in lowercase to compare with the LLM response
    relevant_answers = ["yes", "yes."]
    if text.lower() in relevant_answers:
        return True
    else:
        # Print the other answers for debugging purposes
        print(image_path + ": " + text)
        return False

def process_docx(docx_path, output_folder, llm_address, update):
    output_file_path = os.path.join(output_folder, "diagrams.docx")
    output_folder = os.path.join(output_folder, "images")

    with open('prompts.json', 'r') as f:
        prompts = json.load(f)

    if not os.path.exists(docx_path):
        update("Error: the input document was not found.")
        return None
    
    input_doc = docx2python(docx_path, output_folder, html=True)
    output_doc = Document()
    current_section = "No section"

    lines = input_doc.text.splitlines()
    i = 0
    for line in lines:
        update(f"Processing the document...", i / len(lines))
        i += 1
        if "<h2>" in line or "<h3>" in line:
            current_section = line
        if "media/image" in line:
            img_name = line[10:].split('-')[0]
            img_path = os.path.join(output_folder, img_name)

            new_name = img_name.replace(".emf", ".png")
            new_path = os.path.join(output_folder, new_name)
            try:
                Image.open(img_path).save(output_folder + "/" + new_name)
                os.remove(img_path)
                if is_sequence_diagram(new_path, llm_address, prompts):
                    output_doc.add_heading(current_section[4:-4], level=1)
                    output_doc.add_picture(new_path, width=Inches(6))
                    output_doc.add_paragraph(process_sequence_diagram(new_path))
                    output_doc.add_page_break()
            except Exception as e:
                update(f"Error adding image {img_name} to the document.")

    output_doc.save(output_file_path)
    update(f"Finished processing the document. Saved to {output_file_path}")


if __name__ == "__main__":
    process_docx("test_files/23502-i20_l.docx", "output", "localhost:11435", None)
    main()